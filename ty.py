import os
import re
import unicodedata
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
import groq  # make sure you have your groq client installed

# --------------------------------------------------
# Mapping for check marks inserted via symbol fonts (e.g., Wingdings)
# --------------------------------------------------
# In many Word documents, a check mark may be represented by a character like "ü"
# when using a symbol font such as Wingdings. Extend this mapping if needed.
WINGDINGS_MAP = {
    "ü": "✓",  # common representation in Wingdings
    # You can add more mappings if your document uses other characters.
}

# --------------------------------------------------
# Helper: Normalize check marks using Unicode name
# --------------------------------------------------
def normalize_check_marks(text):
    """
    Replaces any character in the text whose Unicode name includes "CHECK MARK"
    with the standard check mark (✓).
    """
    normalized_chars = []
    for ch in text:
        try:
            if "CHECK MARK" in unicodedata.name(ch):
                normalized_chars.append("✓")
            else:
                normalized_chars.append(ch)
        except ValueError:
            # For characters without a Unicode name, just keep them.
            normalized_chars.append(ch)
    return ''.join(normalized_chars)

# --------------------------------------------------
# Helper: Process a paragraph by iterating over its runs.
# This function checks if a run uses a symbol font (e.g., Wingdings)
# and applies a mapping so that check marks appear correctly.
# --------------------------------------------------
def get_normalized_text_from_paragraph(paragraph):
    parts = []
    for run in paragraph.runs:
        text = run.text
        font_name = run.font.name  # may be None if not explicitly set
        font_name_lower = font_name.lower() if font_name else ""
        # If the run is in a known symbol font, use the mapping
        if font_name_lower in ["wingdings", "wingdings 2", "wingdings 3"]:
            mapped_text = "".join(WINGDINGS_MAP.get(ch, ch) for ch in text)
            parts.append(mapped_text)
        else:
            # Otherwise, use our normal Unicode check mark normalization
            parts.append(normalize_check_marks(text))
    return "".join(parts)

# --------------------------------------------------
# Helper: Get normalized text from a table cell.
# --------------------------------------------------
def get_normalized_text_from_cell(cell):
    paragraphs = cell.paragraphs
    return "\n".join(get_normalized_text_from_paragraph(p) for p in paragraphs)

# --------------------------------------------------
# Helper: Iterate over all blocks (paragraphs and tables) in order.
# --------------------------------------------------
def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)

# --------------------------------------------------
# Helper: Interpret a table from a DOCX file.
# The function now uses our normalized text extraction for each cell.
# --------------------------------------------------
def interpret_table(table):
    if not table.rows:
        return ""
    
    num_cols = len(table.rows[0].cells)
    result_lines = ["جدول استخراج شده:"]
    
    if num_cols == 2:
        for row in table.rows:
            key = get_normalized_text_from_cell(row.cells[0]).strip()
            value = get_normalized_text_from_cell(row.cells[1]).strip()
            if key and value:
                result_lines.append(f"برچسب: {key} -> مقدار: {value}")
            elif key:
                result_lines.append(f"برچسب: {key}")
            elif value:
                result_lines.append(f"مقدار: {value}")
        if len(result_lines) == 1:
            return "بدون پاسخ"
    elif len(table.rows) >= 2:
        headers = []
        for idx, cell in enumerate(table.rows[0].cells):
            header_text = get_normalized_text_from_cell(cell).strip()
            if header_text:
                headers.append((idx, header_text))
        if headers:
            header_line = " | ".join([h for _, h in headers])
            result_lines.append("سرصفحه: " + header_line)
        for row in table.rows[1:]:
            row_pairs = []
            for idx, header in headers:
                cell_text = get_normalized_text_from_cell(row.cells[idx]).strip()
                if cell_text:
                    row_pairs.append(f"{header}: {cell_text}")
            if row_pairs:
                result_lines.append("ردیف: " + ", ".join(row_pairs))
        if len(result_lines) <= 2:
            return "بدون پاسخ"
    else:
        row_data = [get_normalized_text_from_cell(cell).strip() for cell in table.rows[0].cells if get_normalized_text_from_cell(cell).strip()]
        if row_data:
            result_lines.append(" | ".join(row_data))
        else:
            return "بدون پاسخ"
    
    return "\n".join(result_lines)

# --------------------------------------------------
# Extract text (paragraphs and tables) from DOCX file.
# Now paragraphs use our per-run normalization.
# --------------------------------------------------
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    lines = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = get_normalized_text_from_paragraph(block).strip()
            if text:
                lines.append(text)
        elif isinstance(block, Table):
            table_text = interpret_table(block)
            if table_text and table_text != "بدون پاسخ":
                lines.append(table_text)
    return "\n".join(lines)

# --------------------------------------------------
# Extract questions from the full text using a regex pattern.
# (Assumes questions are numbered like "1- ..." at the start of a line.)
# --------------------------------------------------
def extract_questions(full_text):
    pattern = re.compile(r'(?ms)^(?P<qnum>\d+-)(?P<qtext>.*?)(?=^\d+-|\Z)')
    questions = []
    for match in pattern.finditer(full_text):
        qnum = match.group("qnum").strip()
        qtext = match.group("qtext").strip()
        full_question = qnum + " " + qtext
        questions.append((qnum, full_question))
    return questions

# --------------------------------------------------
# Get LLM response (streamed response).
# --------------------------------------------------
def get_llm_response(prompt, client):
    try:
        completion = client.chat.completions.create(
            model="qwen-2.5-32b",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_completion_tokens=4096,
            top_p=0.95,
            stream=True,
            stop=None,
        )
        response_text = ""
        for chunk in completion:
            response_text += chunk.choices[0].delta.content or ""
        return response_text.strip()
    except Exception as e:
        print("LLM Error:", e)
        return None

# --------------------------------------------------
# Process each question: send an individual prompt to the LLM and collect results as Q&A pairs.
# --------------------------------------------------
def process_questions(full_text, client):
    questions = extract_questions(full_text)
    if not questions:
        return []
    
    qa_pairs = []
    for qnum, question in questions:
        prompt = f"""You are an AI specialized in text analysis and structured data extraction from Persian survey tables.
Your task is to extract the answer for the following question in a very short, concise, and direct manner.
Instructions:
1. If an answer is provided, return only the answer text.
2. If no answer is provided, return "پاسخ داده نشده".
3. Do not include any extra commentary or formatting.
4. If you see any check marks the provided content, the colsest answer will be the correct answer, and consider that as the correct answer
Here is the text for analysis:
{question}
Answer:"""
        answer = get_llm_response(prompt, client)
        if answer:
            answer = normalize_check_marks(answer)
        else:
            answer = "پاسخ داده نشده"
        processed_question = question[len(qnum):].strip()
        qa_pairs.append((processed_question, answer))
    return qa_pairs

# --------------------------------------------------
# Convert DOCX to a new DOCX file with neatly formatted plain text.
# --------------------------------------------------
def convert_docx_to_docx(input_docx, output_docx, client):
    full_text = extract_text_from_docx(input_docx)
    if not full_text:
        print("❌ هیچ متنی از سند استخراج نشد.")
        return
    
    qa_pairs = process_questions(full_text, client)
    if not qa_pairs:
        print("❌ هیچ سوالی استخراج نشد.")
        return
    
    new_doc = Document()
    new_doc.add_heading("نتایج استخراج سوالات و پاسخ‌ها", level=1)
    
    for question, answer in qa_pairs:
        p_question = new_doc.add_paragraph()
        run_q = p_question.add_run("سوال: ")
        run_q.bold = True
        p_question.add_run(question)
        
        p_answer = new_doc.add_paragraph()
        run_a = p_answer.add_run("پاسخ: ")
        run_a.bold = True
        run_answer = p_answer.add_run(answer)
        run_answer.font.name = 'Tahoma'
        run_answer._element.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')
        
        new_doc.add_paragraph("")
    
    new_doc.save(output_docx)
    print(f"✅ داده‌ها با موفقیت در {output_docx} ذخیره شدند.")

# --------------------------------------------------
# Main Execution
# --------------------------------------------------d
if __name__ == "__main__":
    input_docx = r"C:\Users\i7 11th\Desktop\فرم شناسنامه تحلیلی نهایی شده.docx"
    output_docx = r"C:\Users\i7 11th\Desktop\results.docx"
    
    # Initialize the LLM client with your API key
    client = groq.Groq(api_key="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")
    
    convert_docx_to_docx(input_docx, output_docx, client)
